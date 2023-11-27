# -*- coding: utf-8 -*-
# @Author  : yingzhu
# @Time    : 2023/10/19 20:30
# @File    ：demo2.py
# @Function:


from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

doc2 = Document()  # 新建文档

doc2.add_paragraph('哪个不是水果：')  # 增加无序列表
doc2.add_paragraph('苹果', style='List Bullet')
doc2.add_paragraph('香蕉', style='List Bullet')
doc2.add_paragraph('馄炖', style='List Bullet')

doc2.add_paragraph('2020年度计划：')  # 增加有序列表
doc2.add_paragraph('每周健身一天', style='List Number')
doc2.add_paragraph('学习50本书', style='List Number')
doc2.add_paragraph('减少加班时间', style='List Number')

doc2.add_heading('图片', 2)
doc2.add_picture('python_install.png', width=Inches(5.5))  # 增加图像

doc2.add_heading('表格', 2)
table = doc2.add_table(rows=1, cols=3)  # 增加表格，这是表格头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '编号'
hdr_cells[1].text = '姓名'
hdr_cells[2].text = '职业'
# 这是表格数据
records = (
    (1, '张三', '电工'),
    (2, '张五', '老板'),
    (3, '马六', 'IT')
)
for id, name, work in records:  # 遍历数据并展示
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = name
    row_cells[2].text = work

doc2.add_page_break()  # 手动增加分页

doc2.save('word2.docx')  # 保存文件
