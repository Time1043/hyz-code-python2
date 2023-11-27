# -*- coding: utf-8 -*-
# @Author  : yingzhu
# @Time    : 2023/10/19 20:39
# @File    ：demo3.py
# @Function:


from docx import Document

doc1 = Document('word1.docx')  # 打开文档1
pl = [paragraph.text for paragraph in doc1.paragraphs]  # 读取每段内容
print('###### 输出word1文章内容')
for i in pl:  # 输出读取到的内容
    print(i)

doc2 = Document('word2.docx')  # 打开文档2
pl2 = [paragraph.text for paragraph in doc2.paragraphs]
print('\n###### 输出word2文章内容')
for j in pl2:  # 输出读取到的内容
    print(j)

tables = [table for table in doc2.tables]  # 读取表格材料，并输出结果
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text, end='  ')
        print()
    print('\n')
